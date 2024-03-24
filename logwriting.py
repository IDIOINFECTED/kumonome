from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
import os
import urllib
from urllib.request import urlopen



class Docx_log:
    def __init__(self, title):
        self.title = title
    def default_set(self):
        document = Document()
        style = document.styles['Normal']
        background = OxmlElement('w:background')
        background.set(qn('w:color'), '0D0D0D')
        document.element.insert(0, background)
        background_shp = OxmlElement('w:displayBackgroundShape')
        document.settings.element.insert(0, background_shp)
        style.font.name = 'Calibri'
        style.font.size = Pt(18)
        style.font.color.rgb = RGBColor(158, 16, 28)
        section = document.sections[-1]
        section.top_margin = Mm(9)
        section.bottom_margin = Mm(9)
        section.left_margin = Mm(15)
        section.right_margin = Mm(10)
        document.save(self.title)
    def logwrite(self, line):
        document = Document(self.title)
        p = document.add_paragraph(line)
        document.save(self.title)
    def add_pict(self, picurl, pict):
        document = Document(self.title)
        urllib.request.urlretrieve(picurl, pict)
        document.add_picture(pict, width = Mm(164.2))
        document.save(self.title)
#test = Docx_log("test.docx")
#test.default_set()
#test.logwrite("test")
#test.add_pict('https://i.ytimg.com/vi/EWybN_J1kIw/hqdefault.jpg', 'preview.jpg')
def make_logs_dir(log_mode):
    folder_date = (str(datetime.now().date()))
    folder_name = f"kumonome_logs_{folder_date}"
    try:
        os.mkdir(folder_name)
        if log_mode == 1:
            test_log = Docx_log(f"{folder_name}/logs.docx")
            test_log.default_set()
    except FileExistsError:
        pass
    return folder_name